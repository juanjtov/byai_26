"""
Document Generator service for creating Word documents.

Generates Word documents (.docx) from AI-generated estimates,
styled using organization's format patterns extracted from their uploaded documents.
"""

import io
import re
import logging
from datetime import datetime
from typing import Optional, Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.services.supabase import get_supabase_admin
from app.services.format_extractor import FormatExtractorService

logger = logging.getLogger("remodly.document_generator")


class DocumentGeneratorService:
    """Generate Word documents using organization's format patterns."""

    # Default styling when no org patterns available
    DEFAULT_FONT = "Calibri"
    DEFAULT_HEADING_SIZES = {"h1": 18, "h2": 14, "h3": 12}
    DEFAULT_BODY_SIZE = 11
    DEFAULT_TEXT_COLOR = RGBColor(0, 0, 0)  # Black

    def __init__(self):
        self.admin = get_supabase_admin()
        self.format_extractor = FormatExtractorService()

    def _hex_to_rgb(self, hex_color: str) -> Optional[RGBColor]:
        """Convert hex color string to RGBColor."""
        if not hex_color:
            return None
        try:
            hex_color = hex_color.lstrip('#')
            if len(hex_color) == 6:
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                return RGBColor(r, g, b)
        except (ValueError, TypeError):
            pass
        return None

    def _parse_markdown_line(self, line: str) -> Tuple[str, str, int]:
        """
        Parse a markdown line to determine its type.

        Returns:
            Tuple of (line_type, content, level)
            line_type: 'heading', 'bullet', 'numbered', 'table_row', 'paragraph'
            content: The text content without markdown syntax
            level: For headings (1-3), for lists (indentation level)
        """
        line = line.strip()

        if not line:
            return ('empty', '', 0)

        # Headings: # ## ###
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            content = heading_match.group(2).strip()
            return ('heading', content, level)

        # Bold headings: **Text** or __Text__ at start of line
        bold_heading = re.match(r'^\*\*(.+)\*\*:?\s*$', line)
        if bold_heading:
            return ('heading', bold_heading.group(1), 2)

        # Bullet lists: - or *
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            content = bullet_match.group(2)
            return ('bullet', content, indent)

        # Numbered lists: 1. 2. etc.
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if numbered_match:
            indent = len(numbered_match.group(1)) // 2
            content = numbered_match.group(2)
            return ('numbered', content, indent)

        # Table rows: | col1 | col2 |
        if line.startswith('|') and line.endswith('|'):
            # Skip separator rows (|---|---|)
            if re.match(r'^\|[\s\-:|]+\|$', line):
                return ('table_separator', '', 0)
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            return ('table_row', cells, 0)

        # Regular paragraph
        return ('paragraph', line, 0)

    def _apply_inline_formatting(self, paragraph, text: str, font_name: str, font_size: Pt, font_color: RGBColor):
        """Apply inline markdown formatting (bold, italic) to a paragraph."""
        # Pattern to find **bold** and *italic* text
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|_(.+?)_|`(.+?)`|([^*_`]+))'

        for match in re.finditer(pattern, text):
            full_match = match.group(0)

            run = paragraph.add_run()
            run.font.name = font_name
            run.font.size = font_size
            if font_color:
                run.font.color.rgb = font_color

            if full_match.startswith('**') or full_match.startswith('__'):
                # Bold
                run.text = match.group(2) or match.group(4)
                run.bold = True
            elif full_match.startswith('*') or full_match.startswith('_'):
                # Italic
                run.text = match.group(3) or match.group(5)
                run.italic = True
            elif full_match.startswith('`'):
                # Code/monospace
                run.text = match.group(6)
                run.font.name = "Courier New"
            else:
                # Regular text
                run.text = match.group(7) or full_match

    async def generate_estimate_document(
        self,
        org_id: str,
        content: str,
        title: str = "Estimate"
    ) -> bytes:
        """
        Generate a Word document from estimate content.

        Args:
            org_id: Organization UUID
            content: AI-generated estimate text (markdown format)
            title: Document title

        Returns:
            Document as bytes
        """
        # Get organization's format patterns
        format_patterns = await self.format_extractor.get_org_format_patterns(org_id)

        # Extract styling from patterns or use defaults
        font_name = self.DEFAULT_FONT
        heading_sizes = self.DEFAULT_HEADING_SIZES.copy()
        body_size = self.DEFAULT_BODY_SIZE
        text_color = self.DEFAULT_TEXT_COLOR
        uses_bold = True

        if format_patterns:
            typography = format_patterns.get("typography", {})
            colors = format_patterns.get("colors", {})

            # Font
            if typography.get("primary_font"):
                font_name = typography["primary_font"]

            # Heading sizes
            if typography.get("heading_sizes"):
                for level, size in typography["heading_sizes"].items():
                    if size:
                        heading_sizes[level] = size

            # Colors
            primary_color = colors.get("primary_text_color")
            if primary_color:
                rgb = self._hex_to_rgb(primary_color)
                if rgb:
                    text_color = rgb

            # Emphasis
            uses_bold = typography.get("uses_bold_for_emphasis", True)

        logger.info(f"Generating document with font={font_name}, colors={text_color}")

        # Create document
        doc = Document()

        # Set default font for document
        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(body_size)

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(body_size)
        date_run.font.name = font_name

        doc.add_paragraph()  # Spacing

        # Parse and add content
        lines = content.split('\n')
        current_table_data: List[List[str]] = []
        in_table = False

        for line in lines:
            line_type, line_content, level = self._parse_markdown_line(line)

            # Handle table accumulation
            if line_type == 'table_row':
                if not in_table:
                    in_table = True
                    current_table_data = []
                current_table_data.append(line_content)
                continue
            elif line_type == 'table_separator':
                continue
            elif in_table:
                # End of table, render it
                self._add_table(doc, current_table_data, font_name, body_size, text_color)
                current_table_data = []
                in_table = False

            if line_type == 'empty':
                doc.add_paragraph()

            elif line_type == 'heading':
                heading_level = min(level, 3)
                para = doc.add_heading(line_content, level=heading_level)
                # Apply font to heading
                for run in para.runs:
                    run.font.name = font_name
                    size = heading_sizes.get(f"h{heading_level}", 14)
                    run.font.size = Pt(size)

            elif line_type == 'bullet':
                para = doc.add_paragraph(style='List Bullet')
                self._apply_inline_formatting(
                    para, line_content, font_name, Pt(body_size), text_color
                )

            elif line_type == 'numbered':
                para = doc.add_paragraph(style='List Number')
                self._apply_inline_formatting(
                    para, line_content, font_name, Pt(body_size), text_color
                )

            elif line_type == 'paragraph':
                para = doc.add_paragraph()
                self._apply_inline_formatting(
                    para, line_content, font_name, Pt(body_size), text_color
                )

        # Handle any remaining table
        if in_table and current_table_data:
            self._add_table(doc, current_table_data, font_name, body_size, text_color)

        # Add footer with generation note
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run("Generated by REMODLY AI")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.name = font_name

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info(f"Generated document: {len(buffer.getvalue())} bytes")
        return buffer.getvalue()

    def _add_table(
        self,
        doc: Document,
        table_data: List[List[str]],
        font_name: str,
        font_size: int,
        text_color: RGBColor
    ):
        """Add a table to the document."""
        if not table_data:
            return

        # Determine table dimensions
        num_rows = len(table_data)
        num_cols = max(len(row) for row in table_data) if table_data else 0

        if num_cols == 0:
            return

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(table_data):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    para = cell.paragraphs[0]
                    run = para.add_run(str(cell_text))
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    if text_color:
                        run.font.color.rgb = text_color

                    # Bold first row (header)
                    if row_idx == 0:
                        run.bold = True

    async def upload_document(
        self,
        org_id: str,
        doc_bytes: bytes,
        filename: str
    ) -> str:
        """
        Upload generated document to Supabase Storage.

        Args:
            org_id: Organization UUID
            doc_bytes: Document content as bytes
            filename: Filename for the document

        Returns:
            Signed download URL
        """
        # Create file path in exports subfolder
        file_path = f"{org_id}/exports/{filename}"

        # Upload to storage
        self.admin.storage.from_("documents").upload(
            file_path,
            doc_bytes,
            file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )

        # Create signed URL (1 hour expiry)
        signed_url = self.admin.storage.from_("documents").create_signed_url(
            file_path, 3600  # 1 hour
        )

        return signed_url.get("signedURL") or signed_url.get("signed_url", "")

    async def export_message(
        self,
        org_id: str,
        message_content: str,
        title: str = "Estimate"
    ) -> str:
        """
        Generate and upload a document from a message.

        Args:
            org_id: Organization UUID
            message_content: The AI message content to export
            title: Document title

        Returns:
            Signed download URL
        """
        # Generate document
        doc_bytes = await self.generate_estimate_document(org_id, message_content, title)

        # Create filename with timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d-%H%M%S")
        filename = f"estimate-{timestamp}.docx"

        # Upload and get URL
        download_url = await self.upload_document(org_id, doc_bytes, filename)

        return download_url
