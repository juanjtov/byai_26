"""
Document processor service for text extraction and embedding.

Extracts text from uploaded documents and creates embeddings for RAG.
Also extracts formatting metadata (fonts, colors, styles) for format learning.
Extracts structured pricing data from estimates and addendums.
"""

import io
import logging
import httpx
from pypdf import PdfReader
from typing import Optional, Dict, Tuple, Any
from collections import Counter

logger = logging.getLogger("remodly.document_processor")

from app.services.supabase import get_supabase_admin
from app.services.embedding import EmbeddingService
from app.services.document import DocumentService
from app.services.format_extractor import FormatExtractorService
from app.services.pricing_extractor import PricingExtractorService


class DocumentProcessorService:
    """Service for processing documents and creating embeddings."""

    # Document types that should have format patterns extracted
    FORMAT_EXTRACTABLE_TYPES = ["contract", "estimate", "proposal", "invoice", "quote", "addendum"]

    # Document types that should have pricing data extracted
    PRICING_EXTRACTABLE_TYPES = ["estimate", "addendum", "invoice", "quote", "proposal"]

    def __init__(self):
        self.admin = get_supabase_admin()
        self.embedding_service = EmbeddingService()
        self.doc_service = DocumentService()
        self.format_extractor = FormatExtractorService()
        self.pricing_extractor = PricingExtractorService()

    async def process_and_embed_document(self, doc_id: str, org_id: str) -> None:
        """
        Extract text from document and create embeddings.

        This is designed to run as a background task after document upload.

        Args:
            doc_id: Document UUID
            org_id: Organization UUID
        """
        try:
            logger.info(f"Starting processing for document {doc_id}")

            # Update status to processing
            await self.doc_service.update_document_status(doc_id, "processing")

            # Get document record
            doc = await self.doc_service.get_document(doc_id)
            if not doc:
                logger.error(f"Document {doc_id} not found during processing")
                await self.doc_service.update_document_status(
                    doc_id, "error", {"error": "Document not found during processing"}
                )
                return

            # Download file from storage
            download_url = await self.doc_service.get_download_url(doc["file_path"])
            if not download_url:
                logger.warning(f"Could not get download URL for document {doc_id}")
                await self.doc_service.update_document_status(
                    doc_id, "error", {"error": "Could not get download URL"}
                )
                return

            # Extract text and formatting metadata based on mime type
            mime_type = doc.get("mime_type", "")
            logger.debug(f"Extracting text from document {doc_id}, mime_type={mime_type}")
            text, formatting_metadata = await self.extract_text_with_formatting(download_url, mime_type)

            if text:
                # Create embeddings with section-aware chunking
                chunks_created = await self.embedding_service.embed_document(
                    doc_id,
                    org_id,
                    text,
                    {
                        "name": doc.get("name", ""),
                        "type": doc.get("type", ""),
                    },
                    use_section_chunking=True  # Enable section-aware chunking
                )
                logger.info(f"Document {doc_id}: created {chunks_created} embedding chunks")

                # Extract format patterns for relevant document types
                doc_type = doc.get("type", "").lower()
                format_extracted = False
                if doc_type in self.FORMAT_EXTRACTABLE_TYPES:
                    patterns = await self.format_extractor.extract_format_from_document(
                        doc_id, org_id, text,
                        {"name": doc.get("name", ""), "type": doc_type},
                        formatting_metadata=formatting_metadata
                    )
                    format_extracted = patterns is not None

                # Extract structured pricing for relevant document types
                pricing_extracted = False
                if doc_type in self.PRICING_EXTRACTABLE_TYPES:
                    pricing_data = await self.pricing_extractor.extract_pricing_from_document(
                        doc_id, org_id, text,
                        doc_name=doc.get("name", ""),
                        doc_type=doc_type
                    )
                    pricing_extracted = pricing_data is not None
                    if pricing_extracted:
                        logger.info(
                            f"Document {doc_id}: extracted pricing data, "
                            f"confidence: {pricing_data.get('confidence_score', 0):.2f}"
                        )

                # Update with extracted data
                await self.doc_service.update_document_status(
                    doc_id,
                    "processed",
                    {
                        "text_length": len(text),
                        "chunks_created": chunks_created,
                        "format_extracted": format_extracted,
                        "pricing_extracted": pricing_extracted,
                    }
                )
                logger.info(f"Document {doc_id} processed successfully")
            else:
                # No text extracted, still mark as processed
                logger.warning(f"Document {doc_id}: no text extracted")
                await self.doc_service.update_document_status(
                    doc_id,
                    "processed",
                    {"text_length": 0, "note": "No text extracted"}
                )

        except Exception as e:
            logger.error(f"Error processing document {doc_id}: {str(e)}", exc_info=True)
            await self.doc_service.update_document_status(
                doc_id, "error", {"error": str(e)}
            )

    async def extract_text_with_formatting(
        self, file_url: str, mime_type: str
    ) -> Tuple[Optional[str], Dict[str, Any]]:
        """
        Extract text and formatting metadata from various document types.

        Args:
            file_url: Signed download URL
            mime_type: Document MIME type

        Returns:
            Tuple of (extracted_text, formatting_metadata)
        """
        empty_formatting: Dict[str, Any] = {"fonts": [], "colors": [], "text_emphasis": {}}

        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(file_url, timeout=60.0)
                if response.status_code != 200:
                    logger.warning(f"Failed to download file, status={response.status_code}")
                    return None, empty_formatting
                content = response.content

            # PDF extraction with formatting (using PyMuPDF)
            if "pdf" in mime_type.lower():
                return self._extract_pdf_with_formatting(content)

            # Plain text (no formatting)
            if "text" in mime_type.lower():
                return content.decode("utf-8", errors="ignore"), empty_formatting

            # Word documents (.docx)
            if "wordprocessingml" in mime_type.lower() or "docx" in mime_type.lower():
                return self._extract_docx_with_formatting(content)

            # Excel spreadsheets (.xlsx)
            if "spreadsheetml" in mime_type.lower() or "xlsx" in mime_type.lower():
                return self._extract_xlsx_with_formatting(content)

            logger.warning(f"Unsupported mime type: {mime_type}")
            return None, empty_formatting

        except Exception as e:
            logger.error(f"Text extraction failed for {mime_type}: {str(e)}", exc_info=True)
            return None, empty_formatting

    def _extract_pdf_with_formatting(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and formatting from PDF using PyMuPDF.

        Args:
            content: PDF file bytes

        Returns:
            Tuple of (text, formatting_metadata)
        """
        formatting: Dict[str, Any] = {
            "fonts": [],
            "colors": [],
            "text_emphasis": {"bold_count": 0, "italic_count": 0},
            "font_sizes": [],
        }

        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []
            fonts_counter: Counter = Counter()
            colors_set: set = set()
            sizes_counter: Counter = Counter()

            for page in doc:
                # Extract text
                page_text = page.get_text()
                if page_text:
                    text_parts.append(page_text)

                # Extract formatting from text dict
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                for block in blocks:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                font_name = span.get("font", "")
                                if font_name:
                                    fonts_counter[font_name] += 1
                                    # Check for bold/italic in font name
                                    if "bold" in font_name.lower():
                                        formatting["text_emphasis"]["bold_count"] += 1
                                    if "italic" in font_name.lower() or "oblique" in font_name.lower():
                                        formatting["text_emphasis"]["italic_count"] += 1

                                size = span.get("size")
                                if size:
                                    sizes_counter[round(size)] += 1

                                color = span.get("color")
                                if color is not None:
                                    # Convert color int to hex
                                    hex_color = f"#{color:06x}"
                                    colors_set.add(hex_color)

            doc.close()

            # Get most common fonts and sizes
            formatting["fonts"] = [font for font, _ in fonts_counter.most_common(5)]
            formatting["colors"] = list(colors_set)[:10]
            formatting["font_sizes"] = [size for size, _ in sizes_counter.most_common(5)]

            return "\n\n".join(text_parts), formatting

        except ImportError:
            logger.warning("PyMuPDF not installed, falling back to basic PDF extraction")
            # Fallback to pypdf without formatting
            try:
                reader = PdfReader(io.BytesIO(content))
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n\n".join(text_parts), formatting
            except Exception as e:
                logger.warning(f"PDF extraction failed: {str(e)}")
                return "", formatting

        except Exception as e:
            logger.warning(f"PDF extraction with formatting failed: {str(e)}")
            return "", formatting

    def _extract_docx_with_formatting(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and formatting from Word document.

        Args:
            content: DOCX file bytes

        Returns:
            Tuple of (text, formatting_metadata)
        """
        formatting: Dict[str, Any] = {
            "fonts": [],
            "colors": [],
            "text_emphasis": {"bold_count": 0, "italic_count": 0},
            "heading_styles": {},
            "font_sizes": [],
        }

        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document(io.BytesIO(content))
            text_parts = []
            fonts_counter: Counter = Counter()
            colors_set: set = set()
            sizes_counter: Counter = Counter()

            # Extract paragraphs with formatting
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

                    # Check if it's a heading
                    style_name = paragraph.style.name if paragraph.style else ""
                    if style_name.startswith("Heading"):
                        if style_name not in formatting["heading_styles"]:
                            formatting["heading_styles"][style_name] = {
                                "font": None,
                                "size": None,
                                "bold": False,
                            }

                    # Extract run-level formatting
                    for run in paragraph.runs:
                        # Font name
                        if run.font.name:
                            fonts_counter[run.font.name] += 1

                        # Font size
                        if run.font.size:
                            size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / Pt(1)
                            sizes_counter[round(size_pt)] += 1

                        # Bold/Italic
                        if run.font.bold:
                            formatting["text_emphasis"]["bold_count"] += 1
                        if run.font.italic:
                            formatting["text_emphasis"]["italic_count"] += 1

                        # Color
                        if run.font.color and run.font.color.rgb:
                            hex_color = f"#{run.font.color.rgb}"
                            colors_set.add(hex_color)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        text_parts.append(" | ".join(cells))

            # Get most common fonts and sizes
            formatting["fonts"] = [font for font, _ in fonts_counter.most_common(5)]
            formatting["colors"] = list(colors_set)[:10]
            formatting["font_sizes"] = [size for size, _ in sizes_counter.most_common(5)]

            return "\n".join(text_parts), formatting

        except Exception as e:
            logger.warning(f"DOCX extraction with formatting failed: {str(e)}")
            return "", formatting

    def _extract_xlsx_with_formatting(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and formatting from Excel spreadsheet.

        Args:
            content: XLSX file bytes

        Returns:
            Tuple of (text, formatting_metadata)
        """
        formatting: Dict[str, Any] = {
            "fonts": [],
            "colors": [],
            "text_emphasis": {"bold_count": 0, "italic_count": 0},
            "cell_styles": {},
        }

        try:
            from openpyxl import load_workbook

            # Load with formatting (not data_only)
            wb = load_workbook(io.BytesIO(content), data_only=False)
            text_parts = []
            fonts_counter: Counter = Counter()
            colors_set: set = set()

            for sheet in wb.worksheets:
                text_parts.append(f"## Sheet: {sheet.title}")

                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value))

                            # Extract font info
                            if cell.font:
                                if cell.font.name:
                                    fonts_counter[cell.font.name] += 1
                                if cell.font.bold:
                                    formatting["text_emphasis"]["bold_count"] += 1
                                if cell.font.italic:
                                    formatting["text_emphasis"]["italic_count"] += 1
                                if cell.font.color and cell.font.color.rgb:
                                    # openpyxl color can be ARGB (8 chars) or RGB (6 chars)
                                    rgb = str(cell.font.color.rgb)
                                    if len(rgb) == 8:
                                        rgb = rgb[2:]  # Remove alpha
                                    if rgb and rgb != "000000":
                                        colors_set.add(f"#{rgb}")

                    if row_values:
                        text_parts.append(" | ".join(row_values))

            # Get most common fonts
            formatting["fonts"] = [font for font, _ in fonts_counter.most_common(5)]
            formatting["colors"] = list(colors_set)[:10]

            return "\n".join(text_parts), formatting

        except Exception as e:
            logger.warning(f"XLSX extraction with formatting failed: {str(e)}")
            return "", formatting

    async def reprocess_document(self, doc_id: str, org_id: str) -> None:
        """
        Reprocess a document, deleting existing embeddings, format patterns, and pricing first.

        Args:
            doc_id: Document UUID
            org_id: Organization UUID
        """
        logger.info(f"Starting reprocess for document {doc_id}")

        try:
            # Delete existing embeddings, format patterns, and pricing data
            await self.embedding_service.delete_document_embeddings(doc_id)
            logger.debug(f"Deleted existing embeddings for document {doc_id}")

            await self.format_extractor.delete_document_patterns(doc_id)
            logger.debug(f"Deleted existing format patterns for document {doc_id}")

            await self.pricing_extractor.delete_document_pricing(doc_id)
            logger.debug(f"Deleted existing pricing data for document {doc_id}")

            # Process again
            await self.process_and_embed_document(doc_id, org_id)
            logger.info(f"Reprocess completed for document {doc_id}")
        except Exception as e:
            logger.error(f"Reprocess failed for document {doc_id}: {str(e)}", exc_info=True)
            raise
